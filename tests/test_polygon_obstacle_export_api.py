from collections.abc import Generator
from contextlib import contextmanager
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy import text
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.deps import get_db_session
from app.core import runtime
from app.analysis.rule_result import AnalysisRuleResult
from app.db.base import Base
from app.main import app
from app.models.airport import Airport
from app.models.analysis_task import AnalysisTask
from app.models.import_batch import ImportBatch
from app.models.project import Project
from app.models.report_export import ReportExport
from app.models.runway import Runway
from app.models.station import Station
from app.report.export_payload_builder import _flatten_rule_results, _build_summary, _collect_radar_unmet_obstacles, build_export_payload


def _read_valid_excel_bytes() -> bytes:
    return Path("docs/import_demo.xlsx").read_bytes()


class _DispatchRecorder:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def delay(self, task_id: str) -> None:
        self.calls.append({"task_id": task_id})


def _run_import_task() -> None:
    with next(iter(app.dependency_overrides[get_db_session]())) as session:
        from app.application.polygon_obstacle_import import PolygonObstacleImportService

        service = PolygonObstacleImportService(session)
        service.run_import_task("import-batch-1")


def _run_analysis_task(task_id: str) -> None:
    with next(iter(app.dependency_overrides[get_db_session]())) as session:
        from app.application.polygon_obstacle_import import PolygonObstacleImportService

        service = PolygonObstacleImportService(session)
        service.run_analysis_task(task_id)


def _run_export_task(task_id: str) -> None:
    with next(iter(app.dependency_overrides[get_db_session]())) as session:
        from app.application.polygon_obstacle_import import PolygonObstacleImportService

        service = PolygonObstacleImportService(session)
        service.run_export_task(task_id)


@contextmanager
def _create_test_client() -> Generator[TestClient, None, None]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    Base.metadata.create_all(
        bind=engine,
        tables=[
            Project.__table__,
            ImportBatch.__table__,
            AnalysisTask.__table__,
            Airport.__table__,
            Runway.__table__,
            Station.__table__,
            ReportExport.__table__,
        ],
    )
    with engine.begin() as connection:
        connection.execute(
            text(
                """
                CREATE TABLE obstacles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    project_id INTEGER NOT NULL,
                    name VARCHAR(255) NOT NULL,
                    obstacle_type VARCHAR(100),
                    source_batch_id VARCHAR(100),
                    source_row_no INTEGER,
                    top_elevation NUMERIC(10, 2),
                    raw_payload JSON,
                    geom TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP NOT NULL,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP NOT NULL
                )
                """
            )
        )

    def _override_session() -> Generator[Session, None, None]:
        session = testing_session_local()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides = {get_db_session: _override_session}

    with TestClient(app) as client:
        yield client

    Base.metadata.drop_all(
        bind=engine,
        tables=[
            Station.__table__,
            Runway.__table__,
            Airport.__table__,
            ReportExport.__table__,
            AnalysisTask.__table__,
            ImportBatch.__table__,
            Project.__table__,
        ],
    )
    app.dependency_overrides = {}


def _create_succeeded_analysis_task(client: TestClient) -> str:
    app.state.dispatch_import_task = _DispatchRecorder().delay
    runtime.dispatch_import_task = app.state.dispatch_import_task
    app.state.dispatch_analysis_task = _DispatchRecorder().delay
    runtime.dispatch_analysis_task = app.state.dispatch_analysis_task

    create_import_response = client.post(
        "/polygon-obstacle/import",
        data={
            "projectName": "Wuhan Demo",
            "obstacleType": "building",
        },
        files={"excelFile": ("import_demo.xlsx", _read_valid_excel_bytes())},
    )
    import_task_id = create_import_response.json()["taskId"]
    _run_import_task()

    with next(iter(app.dependency_overrides[get_db_session]())) as session:
        session.add(
            Airport(
                name="Airport Near",
                longitude=103.975864,
                latitude=30.506881,
            )
        )
        session.commit()

    create_analysis_response = client.post(
        "/polygon-obstacle/analysis",
        json={
            "importTaskId": import_task_id,
            "targetIds": [1],
        },
    )
    analysis_task_id = create_analysis_response.json()["analysisTaskId"]
    _run_analysis_task(analysis_task_id)
    return analysis_task_id


def test_create_export_task_returns_minimal_task_payload() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        dispatch_recorder = _DispatchRecorder()
        app.state.dispatch_export_task = dispatch_recorder.delay
        runtime.dispatch_export_task = dispatch_recorder.delay

        response = client.post(f"/polygon-obstacle/analysis/{analysis_task_id}/export")

    assert response.status_code == 201
    assert response.json() == {
        "exportTaskId": "export-task-1",
        "analysisTaskId": analysis_task_id,
        "status": "pending",
        "message": "export task created",
        "progressPercent": 0,
    }
    assert dispatch_recorder.calls == [{"task_id": "export-task-1"}]


def test_create_export_task_returns_404_for_unknown_analysis_task() -> None:
    with _create_test_client() as client:
        response = client.post("/polygon-obstacle/analysis/missing-task/export")

    assert response.status_code == 404
    assert response.json() == {"detail": "analysis task not found"}


def test_create_export_task_returns_409_for_pending_analysis_task() -> None:
    with _create_test_client() as client:
        app.state.dispatch_import_task = _DispatchRecorder().delay
        runtime.dispatch_import_task = app.state.dispatch_import_task
        app.state.dispatch_analysis_task = _DispatchRecorder().delay
        runtime.dispatch_analysis_task = app.state.dispatch_analysis_task
        create_import_response = client.post(
            "/polygon-obstacle/import",
            data={
                "projectName": "Wuhan Demo",
                "obstacleType": "building",
            },
            files={"excelFile": ("import_demo.xlsx", _read_valid_excel_bytes())},
        )
        import_task_id = create_import_response.json()["taskId"]

        with next(iter(app.dependency_overrides[get_db_session]())) as session:
            session.add(
                Airport(
                    name="Airport Near",
                    longitude=103.975864,
                    latitude=30.506881,
                )
            )
            session.commit()

        create_analysis_response = client.post(
            "/polygon-obstacle/analysis",
            json={
                "importTaskId": import_task_id,
                "targetIds": [1],
            },
        )
        analysis_task_id = create_analysis_response.json()["analysisTaskId"]

        response = client.post(f"/polygon-obstacle/analysis/{analysis_task_id}/export")

    assert response.status_code == 409
    assert response.json() == {"detail": "analysis task is not ready for export"}


def test_get_export_task_status_returns_existing_task() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/{export_task_id}/status"
        )

    assert response.status_code == 200
    assert response.json() == {
        "exportTaskId": export_task_id,
        "analysisTaskId": analysis_task_id,
        "status": "pending",
        "message": "export task created",
        "progressPercent": 0,
    }


def test_get_export_task_result_returns_empty_payload_before_completion() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/{export_task_id}/result"
        )

    assert response.status_code == 200
    assert response.json() == {
        "exportTaskId": export_task_id,
        "analysisTaskId": analysis_task_id,
        "status": "pending",
        "fileName": None,
        "downloadUrl": None,
        "errorMessage": None,
    }


def test_get_export_task_status_returns_404_for_unknown_task() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/missing-task/status"
        )

    assert response.status_code == 404
    assert response.json() == {"detail": "export task not found"}


def test_get_export_task_result_returns_404_for_unknown_task() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/missing-task/result"
        )

    assert response.status_code == 404
    assert response.json() == {"detail": "export task not found"}


def test_download_export_file_returns_404_before_generation() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        response = client.get(f"/polygon-obstacle/exports/{export_task_id}/download")

    assert response.status_code == 404
    assert response.json() == {"detail": "export file not found"}


def test_run_export_task_marks_status_succeeded_and_returns_download_url() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        _run_export_task(export_task_id)
        status_response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/{export_task_id}/status"
        )
        result_response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/{export_task_id}/result"
        )

    assert status_response.status_code == 200
    assert status_response.json() == {
        "exportTaskId": export_task_id,
        "analysisTaskId": analysis_task_id,
        "status": "succeeded",
        "message": "export task succeeded",
        "progressPercent": 100,
    }
    assert result_response.status_code == 200
    assert result_response.json() == {
        "exportTaskId": export_task_id,
        "analysisTaskId": analysis_task_id,
        "status": "succeeded",
        "fileName": f"polygon-obstacle-analysis-{analysis_task_id}.docx",
        "downloadUrl": f"/polygon-obstacle/exports/{export_task_id}/download",
        "errorMessage": None,
    }


def test_build_export_payload_includes_rule_results_with_standards() -> None:
    analysis_task = AnalysisTask(
        id="analysis-task-1",
        import_batch_id="import-batch-1",
        selected_target_ids=[1],
        status="succeeded",
        result_payload={
            "summary": "done",
            "obstacleCount": 1,
            "selectedTargets": [{"id": 1, "name": "Airport A", "category": "机场"}],
            "ruleResults": [
                {
                    "airportId": 1,
                    "stationId": 101,
                    "stationName": "NDB Station",
                    "stationType": "NDB",
                    "obstacleId": 2,
                    "obstacleName": "Obstacle A",
                    "rawObstacleType": "建筑物/构建物",
                    "globalObstacleCategory": "building_general",
                    "ruleCode": "ndb_minimum_distance_50m",
                    "ruleName": "ndb_minimum_distance_50m",
                    "zoneCode": "ndb_minimum_distance_50m",
                    "zoneName": "NDB 50米最小间距",
                    "regionCode": "default",
                    "regionName": "default",
                    "isApplicable": True,
                    "isCompliant": False,
                    "message": "在50米以内",
                    "metrics": {"minimumDistanceMeters": 50.0},
                    "standards": {
                        "gb": {
                            "code": "GB_NDB_50m最小间距区域_50",
                            "text": "GB text",
                            "isCompliant": False,
                        },
                        "mh": {
                            "code": "MH_NDB_50m最小间距区域_50",
                            "text": "MH text",
                            "isCompliant": False,
                        },
                    },
                }
            ],
        },
    )

    payload = build_export_payload(analysis_task)

    assert "tableRows" in payload
    assert len(payload["tableRows"]) == 2  # gb + mh
    row = payload["tableRows"][0]
    assert row["obstacleName"] == "Obstacle A"
    assert row["obstacleType"] == "建筑物/构建物"
    assert row["stationName"] == "NDB Station"
    assert row["standardName"] == "《GB6364-2013》"
    # second row is MH standard
    row_mh = payload["tableRows"][1]
    assert row_mh["standardName"] == "《MH/T 4003.1-2021》"
    assert row_mh["stationName"] == "NDB Station"


def test_download_export_file_returns_docx_after_generation() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        _run_export_task(export_task_id)
        response = client.get(f"/polygon-obstacle/exports/{export_task_id}/download")

    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert (
        f'filename="polygon-obstacle-analysis-{analysis_task_id}.docx"'
        in response.headers["content-disposition"]
    )
    assert len(response.content) > 0


def test_run_export_task_writes_gb_and_mh_standards_into_report() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)

        with next(iter(app.dependency_overrides[get_db_session]())) as session:
            task = session.get(AnalysisTask, analysis_task_id)
            assert task is not None
            task.status = "succeeded"
            task.status_message = "analysis task succeeded"
            task.error_message = None
            task.result_payload = {
                "summary": "已完成测试分析",
                "obstacleCount": 1,
                "selectedTargets": [
                    {"id": 1, "name": "Airport Near", "category": "机场"}
                ],
                "ruleResults": [
                    {
                        "airportId": 1,
                        "stationId": 101,
                        "stationName": "NDB Station",
                        "stationType": "NDB",
                        "obstacleId": 2,
                        "obstacleName": "Obstacle A",
                        "rawObstacleType": "建筑物/构建物",
                        "globalObstacleCategory": "building_general",
                        "ruleCode": "ndb_minimum_distance_50m",
                        "ruleName": "ndb_minimum_distance_50m",
                        "zoneCode": "ndb_minimum_distance_50m",
                        "zoneName": "NDB 50米最小间距",
                        "regionCode": "default",
                        "regionName": "default",
                        "isApplicable": True,
                        "isCompliant": False,
                        "message": "在50米以内",
                        "metrics": {"minimumDistanceMeters": 50.0},
                        "standards": {
                            "gb": {
                                "code": "GB_NDB_50m最小间距区域_50",
                                "text": "GB text",
                                "isCompliant": False,
                            },
                            "mh": {
                                "code": "MH_NDB_50m最小间距区域_50",
                                "text": "MH text",
                                "isCompliant": False,
                            },
                        },
                    }
                ],
            }
            session.commit()

        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]
        _run_export_task(export_task_id)

        with next(iter(app.dependency_overrides[get_db_session]())) as session:
            report_export = session.get(ReportExport, export_task_id)
            assert report_export is not None
            document = Document(report_export.file_path)

    # Check paragraphs for metadata (standards used, etc.)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "GB6364" in full_text
    assert "MH/T 4003.1-2021" in full_text

    # Check table cells for data row content
    all_cell_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Obstacle A" in all_cell_text
    assert "建筑物/构建物" in all_cell_text
    assert "NDB Station" in all_cell_text
    assert "GB text" in all_cell_text
    assert "MH text" in all_cell_text
    assert "《GB6364-2013》" in all_cell_text
    assert "《MH/T 4003.1-2021》" in all_cell_text


def test_run_export_task_succeeds_when_runtime_settings_are_uninitialized() -> None:
    with _create_test_client() as client:
        analysis_task_id = _create_succeeded_analysis_task(client)
        app.state.dispatch_export_task = _DispatchRecorder().delay
        runtime.dispatch_export_task = app.state.dispatch_export_task
        create_response = client.post(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export"
        )
        export_task_id = create_response.json()["exportTaskId"]

        original_settings = runtime.settings
        runtime.settings = None
        try:
            _run_export_task(export_task_id)
        finally:
            runtime.settings = original_settings

        response = client.get(
            f"/polygon-obstacle/analysis/{analysis_task_id}/export/{export_task_id}/result"
        )

    assert response.status_code == 200
    assert response.json()["status"] == "succeeded"
    assert response.json()["downloadUrl"] == (
        f"/polygon-obstacle/exports/{export_task_id}/download"
    )
    assert response.json()["errorMessage"] is None


def test_build_rule_result_payload_includes_is_mid_is_filter_fields() -> None:
    with _create_test_client() as _:
        with next(iter(app.dependency_overrides[get_db_session]())) as session:
            from app.application.polygon_obstacle_import import PolygonObstacleImportService

            service = PolygonObstacleImportService(session)

            result = AnalysisRuleResult(
                station_id=1,
                station_type="NDB",
                obstacle_id=2,
                obstacle_name="Obstacle A",
                raw_obstacle_type="建筑物/构建物",
                global_obstacle_category="building_general",
                rule_code="ndb_minimum_distance_50m",
                rule_name="ndb_minimum_distance_50m",
                zone_code="ndb_minimum_distance_50m",
                zone_name="NDB 50m minimum distance zone",
                region_code="default",
                region_name="default",
                is_applicable=True,
                is_compliant=False,
                message="位于NDB 50米最小间距区域内",
                metrics={"actualDistanceMeters": 30.0, "minimumDistanceMeters": 50.0},
                is_mid=True,
                is_filter_limit=True,
                is_filter_intersect=False,
                standards_rule_code="ndb_minimum_distance_50m",
            )

            payload = service._build_rule_result_payload(
                result=result,
                station_name="NDB Station",
            )

            assert payload["isMid"] is True
            assert payload["isFilterLimit"] is True
            assert payload["isFilterIntersect"] is False


def test_build_rule_result_payload_defaults_is_mid_fields_to_false() -> None:
    with _create_test_client() as _:
        with next(iter(app.dependency_overrides[get_db_session]())) as session:
            from app.application.polygon_obstacle_import import PolygonObstacleImportService

            service = PolygonObstacleImportService(session)

            result = AnalysisRuleResult(
                station_id=1,
                station_type="NDB",
                obstacle_id=2,
                obstacle_name="Obstacle A",
                raw_obstacle_type="建筑物/构建物",
                global_obstacle_category="building_general",
                rule_code="ndb_minimum_distance_50m",
                rule_name="ndb_minimum_distance_50m",
                zone_code="ndb_minimum_distance_50m",
                zone_name="NDB 50m minimum distance zone",
                region_code="default",
                region_name="default",
                is_applicable=True,
                is_compliant=False,
                message="位于NDB 50米最小间距区域内",
                metrics={"actualDistanceMeters": 30.0, "minimumDistanceMeters": 50.0},
            )

            payload = service._build_rule_result_payload(
                result=result,
                station_name="NDB Station",
            )

            assert payload["isMid"] is False
            assert payload["isFilterLimit"] is False
            assert payload["isFilterIntersect"] is False


class TestFlattenRuleResultsSpecialDisplay:
    """T1-T6: Special display logic in _flatten_rule_results."""

    def _make_rule(self, **overrides):
        base = {
            "isApplicable": True,
            "isCompliant": True,
            "zoneCode": "ndb_minimum_distance_50m",
            "ruleCode": "ndb_minimum_distance_50m",
            "obstacleName": "Test Obstacle",
            "rawObstacleType": "building_general",
            "stationName": "Test Station",
            "metrics": {"allowedHeightMeters": 100.0, "overHeightMeters": 5.0},
            "standards": {
                "gb": [{"code": "GB_TEST", "text": "test clause"}],
                "mh": [],
            },
            "message": "",
        }
        base.update(overrides)
        return base

    # ---- T1: isMid shows "不判断" ----
    def test_is_mid_shows_no_judge(self):
        r = self._make_rule(isMid=True, isCompliant=False)
        rows = _flatten_rule_results([r])
        assert len(rows) == 1
        row = rows[0]
        assert row["isCompliant"] is False
        assert row["complianceStatus"] == "不判断"
        assert row["heightLimit"] == "/"
        assert row["overHeight"] == "/"
        assert row["finalOverHeight"] == 0

    # ---- T2: isFilterLimit same as isMid ----
    def test_is_filter_limit_shows_no_judge(self):
        r = self._make_rule(isFilterLimit=True, isCompliant=False)
        rows = _flatten_rule_results([r])
        assert len(rows) == 1
        row = rows[0]
        assert row["isCompliant"] is False
        assert row["complianceStatus"] == "不判断"
        assert row["heightLimit"] == "/"
        assert row["overHeight"] == "/"
        assert row["finalOverHeight"] == 0

    # ---- T3: isFilterIntersect skipped ----
    def test_is_filter_intersect_skipped(self):
        r = self._make_rule(isFilterIntersect=True)
        rows = _flatten_rule_results([r])
        assert len(rows) == 0

    # ---- T4: LOC building restriction special message ----
    def test_loc_building_restriction_special_message(self):
        r = self._make_rule(
            zoneCode="loc_building_restriction_zone",
            metrics={
                "enteredProtectionZone": True,
                "allowedHeightMeters": 80.0,
                "overHeightMeters": 12.5,
            },
        )
        rows = _flatten_rule_results([r])
        assert len(rows) == 1
        row = rows[0]
        assert "建议结合MH4003.1-2021" in row["complianceStatus"]
        assert row["heightLimit"] == 80.0
        assert row["overHeight"] == 12.5
        assert row["finalOverHeight"] == 0

    # ---- T5: Radar 16KM special message ----
    def test_radar_16km_special_message(self):
        r = self._make_rule(
            ruleCode="radar_rotating_reflector_16km",
            zoneCode="radar_rotating_reflector_16km",
            metrics={
                "enteredProtectionZone": True,
                "allowedHeightMeters": 200.0,
                "overHeightMeters": 8.0,
            },
        )
        rows = _flatten_rule_results([r])
        assert len(rows) == 1
        row = rows[0]
        assert "16km范围内" in row["complianceStatus"]
        assert row["heightLimit"] == 200.0
        assert row["overHeight"] == 8.0
        assert row["finalOverHeight"] == 0

    # ---- T6a: isApplicable=False with isMid still shows row ----
    def test_is_applicable_false_with_isMid_still_shows_row(self):
        r = self._make_rule(isApplicable=False, isMid=True, isCompliant=False)
        rows = _flatten_rule_results([r])
        assert len(rows) == 1
        row = rows[0]
        assert row["isCompliant"] is False
        assert row["complianceStatus"] == "不判断"

    # ---- T6b: isApplicable=False without isMid is skipped ----
    def test_is_applicable_false_without_isMid_is_skipped(self):
        r = self._make_rule(isApplicable=False, isMid=False, isCompliant=False)
        rows = _flatten_rule_results([r])
        assert len(rows) == 0

    # ---- T6: isMid not tracked for finalOverHeight ----
    def test_is_mid_not_tracked_in_final_over_height(self):
        normal = self._make_rule(
            obstacleName="Obstacle X",
            stationName="Station Y",
            metrics={
                "allowedHeightMeters": 50.0,
                "overHeightMeters": 10.0,
            },
        )
        mid_rule = self._make_rule(
            isMid=True,
            isCompliant=False,
            obstacleName="Obstacle X",
            stationName="Station Y",
            metrics={
                "allowedHeightMeters": 999.0,
                "overHeightMeters": 99.0,
            },
        )
        rows = _flatten_rule_results([normal, mid_rule])
        assert len(rows) == 2
        for row in rows:
            assert row["finalOverHeight"] == 10.0

    # ---- T7: finalOverHeight aggregated by obstacle, not obstacle+station ----
    def test_final_over_height_agg_by_obstacle(self):
        r1 = self._make_rule(
            obstacleName="Building A",
            stationName="Station X",
            metrics={
                "allowedHeightMeters": 50.0,
                "overHeightMeters": 5.0,
            },
        )
        r2 = self._make_rule(
            obstacleName="Building A",
            stationName="Station Y",
            metrics={
                "allowedHeightMeters": 999.0,
                "overHeightMeters": 15.0,
            },
        )
        r3 = self._make_rule(
            obstacleName="Building B",
            stationName="Station X",
            metrics={
                "allowedHeightMeters": 100.0,
                "overHeightMeters": 3.0,
            },
        )
        rows = _flatten_rule_results([r1, r2, r3])
        for row in rows:
            if row["obstacleName"] == "Building A":
                assert row["finalOverHeight"] == 15.0, (
                    f"Expected 15.0 for Building A, got {row['finalOverHeight']} "
                    f"at station {row['stationName']}"
                )
            elif row["obstacleName"] == "Building B":
                assert row["finalOverHeight"] == 3.0, (
                    f"Expected 3.0 for Building B, got {row['finalOverHeight']}"
                )


class TestBuildSummaryRadarIntegration:
    """T7-T16: _build_summary skip logic and radar conclusion integration."""

    # ---- T7: _build_summary skips isMid rules ----
    def test_build_summary_skips_is_mid(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "BadObs",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 10.0},
                "isMid": True,
            },
        ]
        result = _build_summary(rr, obstacle_count=1, radar_unmet_obstacle_names=set())
        assert "均满足标准限高要求" in result

    # ---- T8: _build_summary skips LOC BRZ special ----
    def test_build_summary_skips_loc_brz_entered(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "BadObs",
                "zoneCode": "loc_building_restriction_zone",
                "metrics": {"allowedHeightMeters": 10.0, "enteredProtectionZone": True},
            },
        ]
        result = _build_summary(rr, obstacle_count=1, radar_unmet_obstacle_names=set())
        assert "均满足标准限高要求" in result

    # ---- T9: _build_summary skips Radar 16KM special ----
    def test_build_summary_skips_radar_16km_entered(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "BadObs",
                "zoneCode": "some_radar_zone",
                "ruleCode": "radar_rotating_reflector_16km",
                "metrics": {"allowedHeightMeters": 10.0, "enteredProtectionZone": True},
            },
        ]
        result = _build_summary(rr, obstacle_count=1, radar_unmet_obstacle_names=set())
        assert "均满足标准限高要求" in result

    # ---- T10: radar unmet (all fail) ----
    def test_build_summary_radar_all_fail(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_a",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 10.0},
            },
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_b",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 8.0},
            },
        ]
        radar = {"obs_a", "obs_b"}
        result = _build_summary(rr, obstacle_count=2, radar_unmet_obstacle_names=radar)
        assert "obs_a、obs_b不满足雷达累计水平遮蔽角的要求" in result
        assert "其余障碍物均满足" not in result

    # ---- T11: radar unmet (partial) ----
    def test_build_summary_radar_partial_unmet(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_a",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 10.0},
            },
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_b",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 100.0},
            },
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_c",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 200.0},
            },
        ]
        radar = {"obs_a"}
        result = _build_summary(rr, obstacle_count=3, radar_unmet_obstacle_names=radar)
        assert "obs_a不满足雷达累计水平遮蔽角的要求" in result
        assert "其余障碍物均满足标准要求，可按报批高度进行审批" in result

    # ---- T12: radar unmet (single obstacle) ----
    def test_build_summary_radar_single_unmet(self):
        rr: list = []
        radar = {"only_obs"}
        result = _build_summary(rr, obstacle_count=1, radar_unmet_obstacle_names=radar)
        assert "only_obs不满足雷达累计水平遮蔽角的要求" in result
        assert "其余障碍物均满足" not in result

    # ---- T13: no radar unmet (all compliant, multi) ----
    def test_build_summary_no_radar_all_compliant_multi(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_a",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 100.0},
            },
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_b",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 200.0},
            },
        ]
        result = _build_summary(rr, obstacle_count=2, radar_unmet_obstacle_names=set())
        assert "均满足标准限高要求" in result
        assert "均满足标准要求，可按报批高度进行审批" in result

    # ---- T14: no radar unmet (all non-compliant) ----
    def test_build_summary_no_radar_all_noncompliant(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_a",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 10.0},
            },
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_b",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 8.0},
            },
        ]
        result = _build_summary(rr, obstacle_count=2, radar_unmet_obstacle_names=set())
        assert "不满足标准限高要求" in result
        assert "满足标准要求，可按报批高度进行审批" not in result

    # ---- T15: no radar unmet (mixed compliance) ----
    def test_build_summary_no_radar_mixed(self):
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_a",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 10.0},
            },
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_b",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 200.0},
            },
            {
                "isApplicable": True,
                "isCompliant": True,
                "obstacleName": "obs_c",
                "zoneCode": "some_zone",
                "metrics": {"allowedHeightMeters": 300.0},
            },
        ]
        result = _build_summary(rr, obstacle_count=3, radar_unmet_obstacle_names=set())
        assert "不满足标准限高要求" in result
        assert "其余障碍物均满足标准要求，可按报批高度进行审批" in result

    # ---- T16: 同一障碍物多条不满足规则时，取 MIN allowedHeightMeters ----
    @pytest.mark.parametrize("heights,expected_in,expected_not_in", [
        ([(50.0, "zone_A"), (30.0, "zone_B")], "30.00", "50.00"),
        ([(30.0, "zone_A"), (50.0, "zone_B")], "30.00", "50.00"),
    ])
    def test_build_summary_multiple_rules_same_obstacle_takes_min_height(self, heights, expected_in, expected_not_in):
        """同一障碍物多条不满足规则时，限高应取各规则中的最小值（与规则顺序无关）"""
        rr = [
            {
                "isApplicable": True,
                "isCompliant": False,
                "obstacleName": "obs_X",
                "zoneCode": zone,
                "metrics": {"allowedHeightMeters": height},
            }
            for height, zone in heights
        ]
        result = _build_summary(rr, obstacle_count=1, radar_unmet_obstacle_names=set())
        assert expected_in in result
        assert expected_not_in not in result

    # ---- T17: _collect_radar_unmet_obstacles ----
    def test_collect_radar_unmet_obstacles(self):
        cumulative_results = [
            {"isCompliant": "满足", "obstacleNames": ["ok_obs"]},
            {"isCompliant": "不满足", "obstacleNames": ["bad_a", "bad_b"]},
            {"isCompliant": "不满足", "obstacleNames": ["bad_a", "bad_c"]},
        ]
        result = _collect_radar_unmet_obstacles(cumulative_results)
        assert result == {"bad_a", "bad_b", "bad_c"}


class TestPrecisionHelpers:
    def test_floor2_rounds_down(self):
        from app.analysis.result_helpers import floor2 as _floor2
        assert _floor2(1.234) == 1.23
        assert _floor2(1.239) == 1.23
        assert _floor2(1.200) == 1.20
        assert _floor2(0.0) == 0.0

    def test_ceil2_rounds_up(self):
        from app.analysis.result_helpers import ceil2 as _ceil2
        assert _ceil2(1.231) == 1.24
        assert _ceil2(1.230) == 1.23
        assert _ceil2(1.200) == 1.20
        assert _ceil2(0.0) == 0.0


class TestEmptyExportPayload:
    def test_build_export_payload_empty_returns_isEmpty_flag(self):
        """空 tableRows 时 build_export_payload 应返回 isEmpty=True 和 emptyMessage"""
        from unittest.mock import MagicMock

        mock_task = MagicMock()
        mock_task.result_payload = {
            "ruleResults": [],
            "obstacleCount": 0,
            "selectedTargets": [{"name": "测试机场"}],
        }
        mock_task.import_batch = None

        result = build_export_payload(mock_task)
        assert result["isEmpty"] is True
        assert "测试机场" in result["emptyMessage"]
        assert "通信、导航、监视" in result["emptyMessage"]
        assert result["tableRows"] == []
        assert result["summary"] == ""

    def test_template_renders_empty_result(self):
        """空结果时模板应正常渲染 emptyMessage"""
        from pathlib import Path
        from docxtpl import DocxTemplate

        template_path = Path("app/report/templates/analysis_report_template.docx")
        doc = DocxTemplate(str(template_path))
        context = {
            "projectName": "测试",
            "airportName": "测试机场",
            "standardsUsed": "",
            "stationNames": "",
            "cumulativeMaskAngleResults": [],
            "electromagneticZoneResult": "",
            "isEmpty": True,
            "emptyMessage": "该项目不位于测试机场通信、导航、监视台站场地保护区内。",
            "obstacleCount": 0,
            "summary": "",
            "tableRows": [],
            "nonCompliantRows": [],
            "compliantRows": [],
        }
        doc.render(context)
        # Verify rendered output contains the empty message and excludes table structure
        body_xml = doc.docx.element.body.xml
        assert "通信、导航、监视台站场地保护区" in body_xml
        # Verify nonCompliantRows template tag is NOT in output (should have been skipped via {% else %})
        assert "nonCompliantRows" not in body_xml
