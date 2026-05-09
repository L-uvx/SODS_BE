from pathlib import Path
from typing import Any

from docx import Document


# 输出标准条文内容块。
def _append_standards_block(
    document: Document,
    label: str,
    standards_value: object,
) -> None:
    if standards_value is None:
        document.add_paragraph(f"{label}: 无")
        return

    if isinstance(standards_value, list):
        for standard in standards_value:
            _append_standard_block(document, label, standard)
        return

    _append_standard_block(document, label, standards_value)


# 输出单条标准条文内容。
def _append_standard_block(
    document: Document,
    label: str,
    standard: dict[str, Any] | None,
) -> None:
    if standard is None:
        document.add_paragraph(f"{label}: 无")
        return
    document.add_paragraph(f"{label}: {standard['code']}")
    document.add_paragraph(str(standard["text"]))
    is_compliant = standard.get("isCompliant")
    if is_compliant is None:
        document.add_paragraph(f"{label}是否满足: 未知")
        return
    document.add_paragraph(
        f"{label}是否满足: {'满足' if is_compliant else '不满足'}"
    )


# 生成分析结果的 Word 报告文件。
def build_analysis_report_docx(payload: dict[str, Any], output_path: Path) -> None:
    document = Document()
    document.add_heading("多边形障碍物分析报告", level=1)
    document.add_paragraph(f"分析任务编号: {payload['analysisTaskId']}")
    document.add_paragraph(f"导入任务编号: {payload['importTaskId']}")
    document.add_paragraph(f"生成时间: {payload['generatedAt']}")
    document.add_paragraph(f"障碍物数量: {payload['obstacleCount']}")
    document.add_paragraph(f"结论摘要: {payload['summary']}")

    document.add_heading("选中目标", level=2)
    selected_targets = payload.get("selectedTargets", [])
    if not selected_targets:
        document.add_paragraph("无")
    else:
        for target in selected_targets:
            document.add_paragraph(
                f"{target['name']} ({target['category']})",
                style="List Bullet",
            )

    document.add_heading("规则结果与标准依据", level=2)
    rule_results = payload.get("ruleResults", [])
    if not rule_results:
        document.add_paragraph("无")
    else:
        for rule_result in rule_results:
            document.add_paragraph(
                (
                    f"{rule_result['stationName']} / {rule_result['obstacleName']} / "
                    f"{rule_result['ruleName']}"
                )
            )
            document.add_paragraph(
                f"是否满足: {'满足' if rule_result['isCompliant'] else '不满足'}"
            )
            document.add_paragraph(f"判定说明: {rule_result['message']}")
            details = rule_result.get("details")
            if details:
                document.add_paragraph(f"详细说明: {details}")
            metrics = rule_result.get("metrics")
            if isinstance(metrics, dict):
                metrics_parts = []
                actual_distance = metrics.get("actualDistance")
                if actual_distance is not None:
                    metrics_parts.append(f"实际距离: {actual_distance}")
                over_distance = metrics.get("overDistance")
                if over_distance is not None:
                    metrics_parts.append(f"超出距离: {over_distance}")
                allowed_height = metrics.get("allowedHeight")
                if allowed_height is not None:
                    metrics_parts.append(f"允许高度: {allowed_height}")
                if metrics_parts:
                    document.add_paragraph(f"指标: {' / '.join(metrics_parts)}")
            standards = rule_result.get("standards", {})
            _append_standards_block(document, "国标条文", standards.get("gb"))
            _append_standards_block(document, "行标条文", standards.get("mh"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
